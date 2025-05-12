# app/core/document_parser.py
import os
import logging
from typing import Dict, Any, List, Union
from pathlib import Path
import mimetypes
import re

# Optional imports for specific file types
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentParser:
    """Parse various document formats and convert to a consistent format"""
    
    def __init__(self):
        self.supported_extensions = {
            '.txt': self._parse_text,
            '.md': self._parse_text,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,
            '.html': self._parse_html,
            '.htm': self._parse_html
        }
        
        # Check available parsers
        if not PDF_AVAILABLE:
            logger.warning("PyPDF2 not installed. PDF parsing will be limited.")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not installed. DOCX parsing will be limited.")
        
    def parse_document(self, file_path: Union[str, Path], file_content=None, filename=None) -> Dict[str, Any]:
        """
        Parse a document and return structured content.
        Can accept either a file path or file content + filename.
        """
        if file_path:
            path = Path(file_path)
            file_ext = path.suffix.lower()
            filename = path.name
        elif file_content and filename:
            file_ext = Path(filename).suffix.lower()
        else:
            raise ValueError("Either file_path or (file_content AND filename) must be provided")
            
        # Check if extension is supported
        if file_ext not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Special handling for .doc files
        if file_ext == '.doc':
            logger.warning(f"Note: .doc files use an older format that may not parse correctly. Consider converting to .docx.")
        
        logger.info(f"Parsing document: {filename}")
        
        try:
            if file_path:
                with open(file_path, 'rb') as f:
                    content = f.read()
            else:
                content = file_content
                
            # Parse according to file type
            parser_func = self.supported_extensions[file_ext]
            parsed_content, metadata = parser_func(content, filename)
            
            # Build streamlined response format without redundant description
            result = {
                "content": parsed_content,
                "metadata": {
                    "title": metadata.get("title", filename),
                    # Only include brief description, not duplicating content
                    "type": "document",
                    "format": file_ext.lstrip('.'),
                    "filename": filename
                }
            }
            
            # Only include page count for PDFs if available
            if metadata.get("pages"):
                result["metadata"]["pages"] = metadata.get("pages")
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {e}")
            raise
    
    def _extract_title_from_text(self, text: str) -> str:
        """Try to extract a meaningful title from text content"""
        lines = text.strip().split('\n')
        if not lines:
            return "Untitled Document"
            
        # Try first non-empty line
        for line in lines:
            line = line.strip()
            if line and len(line) > 3:  # Ensure it's not just a single character or symbol
                # Look for lines that seem like titles (not too long, no common separators)
                if len(line) < 100 and not any(sep in line for sep in [':', ';', '=', '|']):
                    return line
                # If it's a longer line, try to extract the first meaningful segment
                elif len(line) < 200:
                    # Split by common separators and take first segment
                    for sep in [':', '-', '–', '|', ' - ', ' | ']:
                        if sep in line:
                            potential_title = line.split(sep)[0].strip()
                            if potential_title and len(potential_title) > 3:
                                return potential_title
                    
                    # If still too long, truncate with ellipsis
                    if len(line) > 100:
                        return line[:97] + "..."
                    return line
                
        # If no good title found in first 5 lines, use filename-based approach
        return "Untitled Document"
    
    def _parse_text(self, content: bytes, filename: str) -> tuple:
        """Parse plain text files including Markdown"""
        try:
            # Try UTF-8 first
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            # Fall back to other encodings
            try:
                text = content.decode('latin-1')
            except UnicodeDecodeError:
                text = content.decode('cp1252', errors='replace')
        
        # Extract title from first line or filename
        title = self._extract_title_from_text(text) or Path(filename).stem
        
        # Basic text metadata
        metadata = {
            "title": title,
            "description": text[:200] + "..." if len(text) > 200 else text
        }
        
        return text, metadata
    
    def _parse_html(self, content: bytes, filename: str) -> tuple:
        """Parse HTML content"""
        try:
            from bs4 import BeautifulSoup
            
            # Decode content
            try:
                html = content.decode('utf-8')
            except UnicodeDecodeError:
                html = content.decode('latin-1', errors='replace')
            
            # Parse HTML
            soup = BeautifulSoup(html, 'html.parser')
            
            # Remove script and style tags
            for script in soup(["script", "style"]):
                script.decompose()
                
            # Get text content
            text = soup.get_text(separator='\n')
            
            # Clean up text - remove excessive whitespace
            text = re.sub(r'\n+', '\n', text)
            text = re.sub(r' +', ' ', text)
            text = text.strip()
            
            # Extract title
            title = None
            if soup.title and soup.title.string:
                title = soup.title.string.strip()
            else:
                title = self._extract_title_from_text(text) or Path(filename).stem
                
            # Get meta description
            description = ""
            meta_desc = soup.find("meta", attrs={"name": "description"})
            if meta_desc and meta_desc.get("content"):
                description = meta_desc.get("content")
            else:
                description = text[:200] + "..." if len(text) > 200 else text
                
            metadata = {
                "title": title,
                "description": description
            }
            
            return text, metadata
            
        except ImportError:
            logger.warning("BeautifulSoup not installed. Using basic HTML parsing.")
            # Fallback to basic parsing
            from html.parser import HTMLParser
            
            class TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text = []
                def handle_data(self, data):
                    self.text.append(data)
                    
            parser = TextExtractor()
            parser.feed(content.decode('utf-8', errors='replace'))
            text = ' '.join(parser.text)
            
            metadata = {
                "title": Path(filename).stem,
                "description": text[:200] + "..." if len(text) > 200 else text
            }
            
            return text, metadata
    
    def _parse_pdf(self, content: bytes, filename: str) -> tuple:
        """Parse PDF content"""
        if not PDF_AVAILABLE:
            logger.warning("PyPDF2 not installed, using limited PDF parsing.")
            return f"[PDF content from {filename} - install PyPDF2 for full support]", {"title": filename}
            
        from io import BytesIO
        
        text_content = []
        title = Path(filename).stem
        
        try:
            pdf_file = BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)
            
            # Try to get document info
            if reader.metadata:
                if reader.metadata.title:
                    title = reader.metadata.title
            
            # Extract text from each page
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text_content.append(page.extract_text())
                
            # Combine all pages
            full_text = "\n\n".join(text_content)
            
            # If no title in metadata, try to extract from first page
            if title == Path(filename).stem and text_content:
                first_line = text_content[0].strip().split('\n')[0]
                if first_line and len(first_line) < 100:
                    title = first_line
            
            metadata = {
                "title": title,
                "description": full_text[:200] + "..." if len(full_text) > 200 else full_text,
                "pages": len(reader.pages)
            }
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return f"[Error parsing PDF: {str(e)}]", {"title": filename}
    
    def _parse_docx(self, content: bytes, filename: str) -> tuple:
        """Parse DOCX content"""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not installed, using limited DOCX parsing.")
            return f"[DOCX content from {filename} - install python-docx for full support]", {"title": filename}
            
        from io import BytesIO
        
        # Check if this is actually a .doc file incorrectly using .docx extension
        if filename.lower().endswith('.doc'):
            # Extract what we can using basic text extraction
            try:
                try:
                    text = content.decode('utf-8', errors='replace')
                except:
                    text = content.decode('latin-1', errors='replace')
                
                # Clean up the text to remove binary garbage
                text = "".join(char for char in text if char.isprintable() or char in "\n\r\t")
                text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
                
                # Try to extract meaningful content
                cleaned_lines = []
                for line in text.split('\n'):
                    line = line.strip()
                    # Keep only lines that have a reasonable text-to-garbage ratio
                    if line and len(line) > 5 and sum(c.isalpha() for c in line) > len(line) * 0.5:
                        cleaned_lines.append(line)
                
                if cleaned_lines:
                    extracted_text = "\n".join(cleaned_lines)
                    title = self._extract_title_from_text(extracted_text) or Path(filename).stem
                    return extracted_text, {"title": title}
                else:
                    return f"[Could not extract readable text from .doc file: {filename}. Please convert to .docx format.]", {"title": Path(filename).stem}
            except Exception as e:
                logger.error(f"Error extracting text from .doc file: {e}")
                return f"[Error with .doc file: {str(e)}. Please convert to .docx format.]", {"title": Path(filename).stem}
        
        try:
            docx_file = BytesIO(content)
            doc = docx.Document(docx_file)
            
            # Extract text
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Try to extract title
            title = Path(filename).stem
            if doc.paragraphs and doc.paragraphs[0].text:
                title = doc.paragraphs[0].text.strip()
                if len(title) > 100:
                    title = title[:97] + "..."
            
            metadata = {
                "title": title,
                "description": full_text[:200] + "..." if len(full_text) > 200 else full_text
            }
            
            return full_text, metadata
            
        except Exception as e:
            if "File is not a zip file" in str(e):
                logger.error(f"This appears to be an old .doc format, not .docx: {e}")
                return f"[This file appears to be in the older .doc format. Please convert to .docx format for proper parsing.]", {"title": Path(filename).stem}
            else:
                logger.error(f"Error parsing DOCX: {e}")
                return f"[Error parsing DOCX: {str(e)}]", {"title": filename}

    def batch_process_documents(self, files: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Process multiple documents and return structured content"""
        results = {}
        
        for file_info in files:
            try:
                file_content = file_info.get("content")
                filename = file_info.get("filename")
                
                if not file_content or not filename:
                    logger.warning(f"Missing content or filename for file")
                    continue
                    
                parsed = self.parse_document(None, file_content, filename)
                
                # Add to results using filename as key
                results[filename] = parsed
                
            except Exception as e:
                logger.error(f"Error processing file {file_info.get('filename', 'unknown')}: {e}")
                continue
                
        return results